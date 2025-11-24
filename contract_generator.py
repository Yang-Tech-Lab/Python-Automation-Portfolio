from docx import Document
from docx.shared import Pt, RGBColor # 用来设置字体大小和颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH # 用来对齐文本
from datetime import datetime
import os

print("🚀 合同批量生成器启动...")

# 1. 模拟客户数据 (在真实订单中，这些数据通常来自 Excel)
clients = [
    {"name": "Google Inc.", "price": "10,000", "service": "Data Scraping"},
    {"name": "Tesla Motors", "price": "25,000", "service": "PCB Design"},
    {"name": "SpaceX", "price": "50,000", "service": "Full Stack Automation"}
]

# 创建输出文件夹
output_folder = "Generated_Contracts"
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# 2. 开始循环生成
for client in clients:
    print(f"📄 正在为 {client['name']} 生成合同...")
    
    # --- 创建一个空白 Word 文档 ---
    doc = Document()
    
    # --- A. 添加标题 ---
    heading = doc.add_heading('SERVICE AGREEMENT', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中
    
    # --- B. 添加正文段落 ---
    # f-string 里的内容就是动态替换的
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Client: {client['name']}")
    
    doc.add_paragraph("-" * 30) # 分割线
    
    # 正文内容
    p = doc.add_paragraph("This contract confirms that Yang (The Provider) will provide ")
    p.add_run(f"{client['service']}").bold = True # 加粗服务名称
    p.add_run(" services to the Client.")
    
    p2 = doc.add_paragraph("The total agreed fee for this project is: ")
    run = p2.add_run(f"${client['price']}") 
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0) # 把价格变成绿色，看着吉利
    
    doc.add_paragraph("-" * 30)
    
    # --- C. 添加签字区 ---
    doc.add_paragraph("\n\n") # 空两行
    doc.add_paragraph("Signed by: ____________________")
    doc.add_paragraph("Yang (Developer)")
    
    # --- D. 保存文件 ---
    # 文件名也是自动生成的
    file_name = f"{output_folder}/Contract_{client['name']}.docx"
    doc.save(file_name)
    
    print(f"✅ 已保存: {file_name}")

print("-" * 30)
print(f"🎉 全部完成！请去 [{output_folder}] 文件夹查看你的 Word 文档。")